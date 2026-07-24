"""
Report Generator Script using python-docx
Generates the comprehensive academic report 'Project17_Report.docx' 
incorporating all enhanced JNTUK R23 Project 17 requirements.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def build_academic_report():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title & Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("JNTUK R23 B.TECH INTERDISCIPLINARY PROJECT REPORT\n")
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 41, 59)

    run_sub = title_p.add_run("PROJECT 17: SMART ATTENDANCE & PERFORMANCE ANALYZER\n")
    run_sub.font.size = Pt(18)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(2, 132, 199)

    run_dept = title_p.add_run("Department of Computer Science & Engineering (AI & Data Science)\nAdmitted Cohort 2025 (80 Students: 30 Female, 50 Male) | Academic Year 2025–2026")
    run_dept.font.size = Pt(11)
    run_dept.font.italic = True
    run_dept.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table of Contents Box
    toc_table = doc.add_table(rows=1, cols=1)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = toc_table.cell(0, 0)
    set_cell_background(cell, "F1F5F9")

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("DOCUMENT STRUCTURE & EXECUTIVE SUMMARY\n")
    r.font.bold = True
    r.font.size = Pt(11)

    summary_text = (
        "1. Abstract & Introduction\n"
        "2. Problem Definition, Cohort Scope & Academic Calendar Mapping\n"
        "3. Mathematical Formulations (Dot Products, Credit Weights & Covariance)\n"
        "4. Data Structures (Custom QuickSort Implementation & Derivation)\n"
        "5. System Architecture & Database Layer (MongoDB & CSV Fallback)\n"
        "6. Smart Admin Daily Attendance Logger & Automated Alert Thresholds (<75%)\n"
        "7. Machine Learning Risk Engine (K-Means Clustering & Trajectory Model)\n"
        "8. Visual Analytics & Experimental Results (Gender Backlogs & Trajectories)\n"
        "9. Conclusion & Future Scope\n"
        "10. Examiner Viva Voce Readiness (8 Core Questions & Answers)"
    )
    p.add_run(summary_text).font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # 1. ABSTRACT
    doc.add_heading("1. Abstract", level=1)
    doc.add_paragraph(
        "Educational institutions face significant operational challenges in early identification of academically "
        "at-risk students due to disconnected records between paper registers and internal mark databases. "
        "This report presents 'Smart Attendance & Performance Analyzer', a production-ready interdisciplinary "
        "application developed for the JNTUK R23 curriculum in AI & Data Science (Project 17). The system tracks "
        "an 80-student cohort (30 female, 50 male) across Semester 1-1 and Semester 1-2. It incorporates Linear Algebra "
        "feature vector dot product scoring (S_i · W), credit-weighted evaluation formulas, custom QuickSort from scratch "
        "with median-of-three pivot selection, PyMongo database persistence with an automatic offline CSV fallback engine, "
        "unsupervised K-Means risk clustering (k=3), automated 75% attendance threshold alerts, and Random Forest / Linear Regression trajectory forecasting."
    ).paragraph_format.space_after = Pt(12)

    # 2. SCOPE & ACADEMIC CALENDAR
    doc.add_heading("2. Cohort Scope & Academic Calendar Mapping", level=1)
    doc.add_paragraph(
        "• Cohort Size: 80 Students (30 Female: 25331A0501–25331A0530 | 50 Male: 25331A0531–25331A0580)\n"
        "• Semester 1-1 Timeline: August 4, 2025 – January 21, 2026\n"
        "  Subjects: Linear Algebra & Calculus, C Programming, Engineering Physics, Engineering Graphics, BEEE\n"
        "• Semester 1-2 Timeline: January 26, 2026 – July 9, 2026\n"
        "  Subjects: Differential Equations & Vector Calculus (DEVC), Communicative English, Applied Chemistry, Data Structures, Basic Civil & Mechanical Engineering (BCME)"
    ).paragraph_format.space_after = Pt(12)

    # 3. MATHEMATICAL FORMULATIONS
    doc.add_heading("3. Mathematical Formulations & Credit Weighted Scoring", level=1)
    doc.add_paragraph(
        "1. Student Feature Vector: S_i = [ Attendance_%, Mid1, Mid2, External, Backlog_Count ]\n"
        "2. Composite Score (Dot Product): Score_i = S_i · W = Σ (S_i[j] * W[j]) where W = [0.25, 0.15, 0.15, 0.35, -0.10]\n"
        "3. Credit Weighted Score: Score = (Internal Marks * 0.4) + (Attendance Percentage * 0.6)\n"
        "4. Pearson Correlation Matrix Covariance: r = Σ((X_i - Xbar)(Y_i - Ybar)) / sqrt(Σ(X_i-Xbar)^2 * Σ(Y_i-Ybar)^2)"
    ).paragraph_format.space_after = Pt(12)

    # Worked Example
    doc.add_heading("Worked Numerical Example (Student 25331A0501 - Aaradhya Kotta):", level=2)
    doc.add_paragraph(
        "• Feature Vector: S_1 = [85.0, 22.0, 24.0, 55.0, 0] | Weight Vector: W = [0.25, 0.15, 0.15, 0.35, -0.10]\n"
        "• Term Multiplication:\n"
        "  1. Attendance: 85.0 * 0.25 = 21.250\n"
        "  2. Mid-1: 22.0 * 0.15 = 3.300\n"
        "  3. Mid-2: 24.0 * 0.15 = 3.600\n"
        "  4. External: 55.0 * 0.35 = 19.250\n"
        "  5. Backlog Penalty: 0 * -0.10 = 0.000\n"
        "• Composite Score = 21.250 + 3.300 + 3.600 + 19.250 = 47.40\n"
        "• Credit-Weighted Score = (23.0 * 0.4) + (85.0 * 0.6) = 9.2 + 51.0 = 60.20"
    ).paragraph_format.space_after = Pt(12)

    # 4. DATA STRUCTURES
    doc.add_heading("4. Data Structures (Custom QuickSort Implementation)", level=1)
    doc.add_paragraph(
        "QuickSort is implemented from scratch with median-of-three pivot selection:\n"
        "• Best / Average Case Time Complexity: O(n log n)\n"
        "• Worst Case Time Complexity: O(n^2) (mitigated via median pivot selection)\n"
        "• Auxiliary Space: O(log n) call stack depth"
    ).paragraph_format.space_after = Pt(12)

    # 5. SYSTEM ARCHITECTURE & DATABASE
    doc.add_heading("5. System Architecture & MongoDB Persistence", level=1)
    doc.add_paragraph(
        "The backend uses PyMongo to interact with MongoDB collections (students, attendance_logs, academic_records, ml_results). "
        "If MongoDB is offline or disconnected, the DatabaseManager automatically falls back to local CSV file I/O."
    ).paragraph_format.space_after = Pt(12)

    # 6. SMART ADMIN & ALERTS
    doc.add_heading("6. Smart Admin Attendance & Threshold Alerts (<75%)", level=1)
    doc.add_paragraph(
        "Faculty can log daily classroom attendance by selecting session date, subject, present roll numbers, and absent roll numbers. "
        "The system updates student records and automatically triggers an alert if attendance drops below the 75.0% threshold."
    ).paragraph_format.space_after = Pt(12)

    # 7. MACHINE LEARNING ENGINE
    doc.add_heading("7. Machine Learning Engine (K-Means & Trajectory)", level=1)
    doc.add_paragraph(
        "1. Unsupervised K-Means Clustering (k=3):\n"
        "   - StandardScaler feature normalization\n"
        "   - Centroid sorting to assign: High Achievers (Safe/Green), Moderate Learners (Yellow), At-Risk (Red)\n"
        "   - Validated via Silhouette score and Elbow curve plot\n"
        "2. Trajectory Model: Random Forest Regressor and Linear Regression model predicting Sem 1-2 performance and GPA."
    ).paragraph_format.space_after = Pt(12)

    # 8. VISUAL ANALYTICS
    doc.add_heading("8. Visual Analytics & Experimental Results", level=1)
    doc.add_paragraph(
        "Analytical plots saved in output/charts/:\n"
        "• attendance_vs_marks.png: Attendance vs External Marks correlation\n"
        "• gender_backlog_distribution.png: Gender-wise Backlog Distribution (30 Female vs 50 Male)\n"
        "• growth_trajectory.png: Sem 1-1 -> Sem 1-2 Growth Trajectory Graph\n"
        "• kmeans_clusters.png: K-Means Risk Cluster Scatter Plot\n"
        "• elbow_curve.png: Elbow Method Inertia Plot"
    ).paragraph_format.space_after = Pt(12)

    # 9. CONCLUSION & FUTURE SCOPE
    doc.add_heading("9. Conclusion & Future Scope", level=1)
    doc.add_paragraph(
        "Conclusion: The system delivers a robust interdisciplinary solution aligning Linear Algebra, Data Structures, DB engineering, and ML for JNTUK R23 evaluation.\n"
        "Future Scope: Biometric attendance integration, automated SMS/WhatsApp alerts, and multi-semester LSTM trajectory modeling."
    ).paragraph_format.space_after = Pt(14)

    # 10. VIVA VOCE Q&A
    doc.add_heading("10. Examiner Viva Voce Preparation (Q&A)", level=1)
    viva_qa = [
        ("Q1: Why K-Means clustering and not supervised classification?",
         "Answer: Early semester records lack labeled ground truth target tags. K-Means naturally discovers academic clusters based on geometric feature proximity."),
        ("Q2: Why implement custom QuickSort from scratch?",
         "Answer: Demonstrates core algorithmic mastery, partition mechanics, median-of-three pivot selection, and time complexity derivations required by R23 rubrics."),
        ("Q3: How does vector dot product weighting prevent metric bias?",
         "Answer: Weighs attendance (25%), Mid-1 (15%), Mid-2 (15%), External (35%), and Backlog penalty (-10%) simultaneously rather than relying on unweighted averages."),
        ("Q4: What happens if MongoDB is offline during your viva demo?",
         "Answer: The database manager catches connection timeout errors and automatically falls back to local CSV file I/O without throwing unhandled exceptions."),
        ("Q5: Why is StandardScaler mandatory before K-Means?",
         "Answer: K-Means relies on Euclidean distance. Attendance (0-100%) would dominate backlog count (0-4) by orders of magnitude without feature scaling."),
        ("Q6: How does median-of-three pivot selection prevent O(n^2) worst-case QuickSort time complexity?",
         "Answer: Eliminates extreme pivot selection on already sorted arrays by selecting the median among low, middle, and high elements."),
        ("Q7: What is the physical meaning of the Pearson correlation coefficient between attendance and marks?",
         "Answer: Measures linear relationship (-1 to +1). A high positive value empirically proves that regular attendance directly correlates with higher exam performance."),
        ("Q8: How does the trajectory model project Semester 1-2 performance?",
         "Answer: Fits a Random Forest / Linear Regression model on Semester 1-1 attendance, Mid scores, and backlogs to directional project Semester 1-2 external exam performance and GPA.")
    ]

    for q, a in viva_qa:
        p_q = doc.add_paragraph()
        r_q = p_q.add_run(q)
        r_q.font.bold = True
        r_q.font.color.rgb = RGBColor(2, 132, 199)

        p_a = doc.add_paragraph()
        p_a.add_run(a)
        p_a.paragraph_format.left_indent = Inches(0.2)
        p_a.paragraph_format.space_after = Pt(8)

    os.makedirs("report", exist_ok=True)
    report_path = os.path.join("report", "Project17_Report.docx")
    doc.save(report_path)
    print(f"Academic report generated successfully at: {report_path}")

if __name__ == "__main__":
    build_academic_report()
